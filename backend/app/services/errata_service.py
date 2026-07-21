import base64
import json
import mimetypes
import os
import re
import shutil
import subprocess
import tempfile
from copy import deepcopy
from html.parser import HTMLParser
from pathlib import Path
from typing import Literal

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.parts.numbering import NumberingPart
from docx.text.paragraph import Paragraph
from langchain_core.messages import HumanMessage, SystemMessage
from pydantic import BaseModel, Field
from sqlalchemy import text

from app.agent.nodes.llm_client import (
    call_with_retry_and_fallback,
    coerce_token_count,
    get_llm,
    get_runtime_request_settings,
    solve_image,
)
from app.agent.state import AgentState
from app.agent.nodes.formatter import format_node
from app.agent.nodes.solver import solve_node_sync
from app.core.database import SessionLocal, engine
from app.core.database import DEFAULT_DB_PATH
from app.models.domain import AgentLog, ErrataItem, ErrataJob, Task
from app.services.mineru_ingestion import parse_local_file_with_mineru
from app.services.runtime_config import (
    get_prompt_bundle,
    resolve_fallback_models,
)
from app.services.task_artifacts import latest_task_artifact, persist_task_artifact


ERRATA_ROOT = Path(
    os.getenv("ERRATA_DATA_DIR", str(DEFAULT_DB_PATH.parent / "errata"))
).resolve()
DEFAULT_ANCHORS = (
    "勘误处理建议/应该为：",
    "勘误处理建议/应该为:",
    "应该为：",
    "应该为:",
    "改为：",
    "改为:",
)
# 保留旧常量，兼容已有调用方与测试；新逻辑使用 DEFAULT_ANCHORS。
MARKER = DEFAULT_ANCHORS[0]
MODEL_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp"}
ERRATA_WORKFLOW_VERSION = 2


def _ensure_errata_task_column() -> None:
    """允许服务在非 FastAPI 启动路径下处理旧 SQLite 数据库。"""
    with engine.begin() as conn:
        task_columns = {
            row[1] for row in conn.execute(text("PRAGMA table_info(tasks)")).fetchall()
        }
        task_additions = {
            "workflow_type": "VARCHAR NOT NULL DEFAULT 'standard'",
            "source_kind": "VARCHAR",
            "source_id": "VARCHAR",
            "source_item_id": "VARCHAR",
        }
        for name, sql_type in task_additions.items():
            if task_columns and name not in task_columns:
                conn.execute(text(f"ALTER TABLE tasks ADD COLUMN {name} {sql_type}"))
        columns = {row[1] for row in conn.execute(text("PRAGMA table_info(errata_items)")).fetchall()}
        additions = {
            "task_id": "VARCHAR",
            "source_start_index": "INTEGER",
            "source_end_index": "INTEGER",
            "material_docx_path": "TEXT",
            "material_paths_json": "TEXT",
            "question_material_paths_json": "TEXT",
            "material_text": "TEXT",
            "material_version": "INTEGER NOT NULL DEFAULT 0",
        }
        for name, sql_type in additions.items():
            if columns and name not in columns:
                conn.execute(text(f"ALTER TABLE errata_items ADD COLUMN {name} {sql_type}"))


class ErrataAdjudication(BaseModel):
    standard_answer_verdict: Literal["correct", "incorrect", "insufficient_evidence"]
    question_verdict: Literal["correct", "incorrect", "insufficient_evidence"]
    original_answer_verdict: Literal["correct", "incorrect", "insufficient_evidence"]
    correction_opinion_verdict: Literal[
        "correct", "partial", "incorrect", "not_provided", "insufficient_evidence"
    ]
    errata_opinion: str = Field(default="")
    question_errata: str = Field(default="")
    warnings: list[str] = Field(default_factory=list)


def derive_errata_result_type(decision: ErrataAdjudication) -> str:
    if (
        decision.standard_answer_verdict == "insufficient_evidence"
        or decision.question_verdict == "insufficient_evidence"
        or decision.original_answer_verdict == "insufficient_evidence"
        or decision.correction_opinion_verdict == "insufficient_evidence"
    ):
        return "insufficient_evidence"
    if decision.question_verdict == "incorrect":
        return "question_errata"
    return "correct" if decision.original_answer_verdict == "correct" else "rewrite"


def errata_allow_write(decision: ErrataAdjudication, has_correction_opinion: bool) -> bool:
    return (
        decision.standard_answer_verdict == "correct"
        and decision.question_verdict != "insufficient_evidence"
        and decision.original_answer_verdict != "insufficient_evidence"
        and (
            decision.question_verdict != "incorrect"
            or bool((decision.question_errata or "").strip())
        )
        and (
            not has_correction_opinion
            or decision.correction_opinion_verdict != "insufficient_evidence"
        )
    )


def build_errata_word_content(decision: ErrataAdjudication, formatted_solution: str) -> str:
    solution = (formatted_solution or "").strip()
    if decision.standard_answer_verdict != "correct":
        raise ValueError("标准答案未通过裁决，不能生成最终正文")
    if decision.question_verdict == "incorrect":
        question_errata = (decision.question_errata or "").strip()
        if not question_errata:
            raise ValueError("题干错误但缺少 question_errata")
        if not solution:
            raise ValueError("题干错误时缺少 Formatter 标准答案")
        return f"【题干勘误】{question_errata}\n\n{solution}"
    if decision.original_answer_verdict == "correct":
        return "原答案正确。"
    if decision.original_answer_verdict == "incorrect":
        if not solution:
            raise ValueError("原答案错误时缺少 Formatter 标准答案")
        return solution
    raise ValueError("证据不足，不能生成最终正文")


def _item_workflow_artifacts(task: Task | None) -> tuple[str, dict]:
    if not task:
        return "", {}
    revision = int(task.input_revision or 1)
    solution = latest_task_artifact(task.task_id, "formatter", revision)
    decision = latest_task_artifact(task.task_id, "errata_adjudication", revision)
    try:
        decision_data = json.loads(decision.content) if decision else {}
    except json.JSONDecodeError:
        decision_data = {}
    return (solution.content if solution else ""), decision_data


def item_to_dict(item: ErrataItem, task: Task | None = None) -> dict:
    solution_text, adjudication = _item_workflow_artifacts(task)
    result_type = item.result_type
    allow_write = False
    try:
        decision = ErrataAdjudication.model_validate(adjudication)
        result_type = derive_errata_result_type(decision)
        allow_write = errata_allow_write(
            decision, bool((item.correction_opinion or "").strip())
        )
    except ValueError:
        pass
    return {
        "item_id": item.item_id,
        "job_id": item.job_id,
        "item_index": item.item_index,
        "task_id": item.task_id,
        "source_ref": item.source_ref or "",
        "question_text": item.question_text or "",
        "original_answer": item.original_answer or "",
        "correction_opinion": item.correction_opinion or "",
        "existing_content": item.existing_content or "",
        "evidence": json.loads(item.evidence_json or "[]"),
        "material_paths": json.loads(item.material_paths_json or "[]"),
        "question_material_paths": json.loads(item.question_material_paths_json or "[]"),
        "material_text": item.material_text or "",
        "material_version": int(item.material_version or 0),
        "has_material_packet": bool(item.material_docx_path),
        "status": task.state if task else item.status,
        "task_state": task.state if task else None,
        "human_confirmed": item.status == "confirmed",
        "result_type": result_type,
        "allow_write": allow_write,
        "final_text_markup": (task.final_result if task else item.final_text_markup) or "",
        "warnings": json.loads(item.warnings_json or "[]"),
        "mineru_text": item.mineru_text or "",
        "review_status": item.review_status,
        "review_feedback": item.review_feedback or "",
        "replace_existing": bool(item.replace_existing),
        "solution_text": solution_text,
        "standard_answer_verdict": adjudication.get("standard_answer_verdict", ""),
        "question_verdict": adjudication.get("question_verdict", ""),
        "original_answer_verdict": adjudication.get("original_answer_verdict", ""),
        "correction_opinion_verdict": adjudication.get("correction_opinion_verdict", ""),
        "errata_opinion": adjudication.get("errata_opinion", ""),
        "question_errata": adjudication.get("question_errata", ""),
    }


def _errata_attachment_urls(item: ErrataItem) -> list[str]:
    paths = list(dict.fromkeys(json.loads(item.question_material_paths_json or "[]")))
    return [f"/api/errata/jobs/{item.job_id}/evidence/{path}" for path in paths]


def sync_errata_task(
    db,
    item: ErrataItem,
    job: ErrataJob | None = None,
    *,
    create_missing: bool = False,
    invalidate: bool = False,
) -> Task | None:
    """同步文档上下文到 Task；仅导入和迁移允许创建缺失 Task。"""
    job = job or db.query(ErrataJob).filter(ErrataJob.job_id == item.job_id).first()
    task = (
        db.query(Task).filter(Task.task_id == item.task_id).first()
        if item.task_id
        else None
    )
    if not task and create_missing:
        task_id = item.task_id or f"task_{item.item_id}"
        task = Task(
            task_id=task_id,
            thread_id=task_id,
            image_url="",
            input_revision=1,
            state="manual",
            retry_count=0,
            workflow_type="errata",
            source_kind="errata",
            source_id=item.job_id,
            source_item_id=item.item_id,
        )
        db.add(task)
        item.task_id = task_id
    if not task:
        return None

    attachments = _errata_attachment_urls(item)
    task.workflow_type = "errata"
    task.source_kind = "errata"
    task.source_id = item.job_id
    task.source_item_id = item.item_id
    task.image_url = attachments[0] if attachments else ""
    task.image_urls_json = json.dumps(attachments, ensure_ascii=False)
    task.question_text = (item.question_text or "").strip() or None
    try:
        history = json.loads(task.history or "{}")
    except Exception:
        history = {}
    history.update(
        {
            "workflow_type": "errata",
            "errata_workflow_version": ERRATA_WORKFLOW_VERSION,
            "errata_item_id": item.item_id,
            "source_kind": "errata",
            "source_id": item.job_id,
            "source_title": job.original_filename if job else "勘误任务",
            "source_item_label": item.source_ref or f"题块 {item.item_index}",
            "attachment_urls": attachments,
            "image_urls": attachments,
        }
    )
    task.history = json.dumps(history, ensure_ascii=False)
    if invalidate:
        task.input_revision = int(task.input_revision or 1) + 1
        task.state = "manual"
        task.current_node = None
        task.final_result = None
        task.question_preview = None
        task.answer_preview = None
        task.error_code = None
        item.review_status = "pending"
        item.review_feedback = None
    return task


def ensure_errata_tasks(db, job_id: str, *, create_missing: bool = False) -> None:
    """兼容迁移入口；普通查询不得通过此函数复活已删除 Task。"""
    job = db.query(ErrataJob).filter(ErrataJob.job_id == job_id).first()
    items = db.query(ErrataItem).filter(ErrataItem.job_id == job_id).all()
    for item in items:
        sync_errata_task(db, item, job, create_missing=create_missing)


def migrate_errata_workflow_v2() -> None:
    migration_id = f"errata_workflow_v{ERRATA_WORKFLOW_VERSION}"
    with engine.begin() as conn:
        conn.execute(
            text(
                "CREATE TABLE IF NOT EXISTS app_migrations "
                "(migration_id VARCHAR PRIMARY KEY, applied_at DATETIME DEFAULT CURRENT_TIMESTAMP)"
            )
        )
        if conn.execute(
            text("SELECT 1 FROM app_migrations WHERE migration_id = :migration_id"),
            {"migration_id": migration_id},
        ).first():
            return

    with SessionLocal() as db:
        rows = (
            db.query(ErrataItem, Task)
            .join(Task, Task.task_id == ErrataItem.task_id)
            .filter(Task.workflow_type == "errata")
            .all()
        )
        for item, task in rows:
            task.input_revision = int(task.input_revision or 1) + 1
            task.state = "manual"
            task.current_node = None
            task.retry_count = 0
            task.final_result = None
            task.question_preview = None
            task.answer_preview = None
            task.error_code = "勘误工作流已升级到 v2，请重新运行。"
            try:
                history = json.loads(task.history or "{}")
            except json.JSONDecodeError:
                history = {}
            for key in (
                "draft_solution",
                "formatted_solution",
                "errata_decision",
                "review_decision",
                "review_feedback",
                "failed_node",
            ):
                history.pop(key, None)
            history["errata_workflow_version"] = ERRATA_WORKFLOW_VERSION
            task.history = json.dumps(history, ensure_ascii=False)
            item.status = "pending"
            item.result_type = None
            item.final_text_markup = None
            item.review_status = "pending"
            item.review_feedback = None
            item.warnings_json = "[]"
        db.execute(
            text("INSERT INTO app_migrations (migration_id) VALUES (:migration_id)"),
            {"migration_id": migration_id},
        )
        db.commit()


def job_to_dict(job: ErrataJob, item_count: int = 0) -> dict:
    return {
        "job_id": job.job_id,
        "original_filename": job.original_filename,
        "state": job.state,
        "error_msg": job.error_msg,
        "mineru_status": job.mineru_status,
        "custom_anchors": job.custom_anchors or "",
        "item_count": item_count,
        "has_output": bool(job.output_path and Path(job.output_path).exists()),
        "created_at": job.created_at,
        "updated_at": job.updated_at,
    }


def _anchors(custom_anchors: str | None = None) -> tuple[str, ...]:
    custom = re.split(r"[\n,，;；]+", custom_anchors or "")
    return tuple(dict.fromkeys((*DEFAULT_ANCHORS, *(value.strip() for value in custom if value.strip()))))


def _normalized_anchor(value: str) -> str:
    return re.sub(r"\s+", "", value).replace(":", "：")


def _is_marker(text: str, anchors: tuple[str, ...]) -> bool:
    normalized_text = _normalized_anchor(text)
    return any(_normalized_anchor(anchor) in normalized_text for anchor in anchors)


def _paragraph_text(paragraph: Paragraph) -> str:
    text_tags = {qn("w:t"), qn("w:delText"), qn("m:t")}
    return "".join(
        node.text or "" for node in paragraph._p.iter() if node.tag in text_tags
    ).strip()


def _body_paragraphs(doc: Document) -> list[Paragraph]:
    return [
        Paragraph(element, element.getparent())
        for element in doc.element.body.iter(qn("w:p"))
    ]


def _paragraph_images(
    paragraph: Paragraph, related_parts: dict, target_dir: Path, prefix: str
) -> list[Path]:
    paths: list[Path] = []
    for index, blip in enumerate(paragraph._p.iter(qn("a:blip")), start=1):
        rel_id = blip.get(qn("r:embed"))
        if not rel_id or rel_id not in related_parts:
            continue
        part = related_parts[rel_id]
        suffix = Path(str(part.partname)).suffix or ".png"
        path = target_dir / f"{prefix}_{index}{suffix}"
        path.write_bytes(part.blob)
        paths.append(path)
    return paths


def _field_text(
    texts: list[str], start: int, end: int, label: str, next_labels: tuple[str, ...]
) -> str:
    label_index = next(
        (idx for idx in range(start, end) if texts[idx].startswith(label)), None
    )
    if label_index is None:
        return ""
    first = texts[label_index][len(label) :].strip()
    lines = [first] if first else []
    for idx in range(label_index + 1, end):
        if any(texts[idx].startswith(next_label) for next_label in next_labels):
            break
        if texts[idx]:
            lines.append(texts[idx])
    return "\n".join(lines).strip()


def _looks_like_next_item_metadata(text: str) -> bool:
    value = text.strip()
    return bool(
        re.fullmatch(r"[A-Z0-9]{3,8}", value)
        or re.match(r"^题号\s*[:：]?\s*\S+", value)
        or re.match(r"^\d{6,8}-.*勘误处理人", value)
        or re.match(r"^\(\d{4}-\d{4}.*\)", value)
    )


def _render_evidence(source_path: Path, work_dir: Path) -> list[Path]:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    pdftoppm = shutil.which("pdftoppm")
    if not soffice or not pdftoppm:
        return []
    render_dir = work_dir / "render"
    render_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(render_dir), str(source_path)],
        check=True,
        capture_output=True,
        timeout=180,
    )
    pdf_path = render_dir / f"{source_path.stem}.pdf"
    if not pdf_path.exists():
        return []
    subprocess.run(
        [pdftoppm, "-png", "-r", "150", str(pdf_path), str(render_dir / "page")],
        check=True,
        capture_output=True,
        timeout=180,
    )
    return sorted(render_dir.glob("page-*.png"))


def _marker_pages(source_path: Path, page_images: list[Path], anchors: tuple[str, ...]) -> list[int]:
    pdftotext = shutil.which("pdftotext")
    pdf_path = source_path.parent / "render" / f"{source_path.stem}.pdf"
    if not pdftotext or not pdf_path.exists():
        return []
    pages: list[int] = []
    for page_number in range(1, len(page_images) + 1):
        result = subprocess.run(
            [pdftotext, "-f", str(page_number), "-l", str(page_number), str(pdf_path), "-"],
            capture_output=True,
            timeout=30,
        )
        page_text = result.stdout.decode("utf-8", errors="ignore")
        pages.extend([page_number] * sum(page_text.count(anchor) for anchor in anchors))
    return pages


def _top_level_body_element(paragraph: Paragraph, body) -> object:
    element = paragraph._p
    while element.getparent() is not body:
        element = element.getparent()
    return element


def _build_material_packet(
    source_path: Path,
    source_doc: Document,
    paragraphs: list[Paragraph],
    start: int,
    end: int,
    item_index: int,
    packet_kind: str = "materials",
) -> tuple[Path, list[str], str]:
    """裁剪原 Word 的题块后渲染，避免用 OCR 字段替代原始版面。"""
    source_dir = source_path.parent
    packet_dir = source_dir / packet_kind / f"item_{item_index}"
    packet_dir.mkdir(parents=True, exist_ok=True)
    packet_path = packet_dir / "source.docx"
    shutil.copy2(source_path, packet_path)

    packet = Document(str(packet_path))
    body = packet.element.body
    source_body = source_doc.element.body
    source_children = list(source_body)
    start_element = _top_level_body_element(paragraphs[start], source_body)
    end_element = _top_level_body_element(paragraphs[end], source_body)
    start_index = source_children.index(start_element)
    end_index = source_children.index(end_element)
    nested_paragraphs_to_remove: list[Paragraph] = []
    if packet_kind == "question_materials":
        packet_children = list(body)
        for paragraph_index, paragraph in enumerate(_body_paragraphs(packet)):
            top_level = _top_level_body_element(paragraph, body)
            top_level_index = packet_children.index(top_level)
            if (
                start_index <= top_level_index <= end_index
                and not start <= paragraph_index <= end
            ):
                nested_paragraphs_to_remove.append(paragraph)
    for index, element in reversed(list(enumerate(list(body)))):
        if element.tag == qn("w:sectPr"):
            continue
        if index < start_index or index > end_index:
            body.remove(element)
    for paragraph in reversed(nested_paragraphs_to_remove):
        paragraph_element = paragraph._p
        parent = paragraph_element.getparent()
        parent.remove(paragraph_element)
        if parent.tag == qn("w:tc") and not any(child.tag == qn("w:p") for child in parent):
            parent.append(OxmlElement("w:p"))
    packet.save(str(packet_path))

    rendered = _render_evidence(packet_path, packet_dir)
    paths = [str(path.relative_to(source_dir)) for path in rendered]
    if not paths:
        assets_dir = packet_dir / "assets"
        assets_dir.mkdir(exist_ok=True)
        for paragraph_index in range(start, end + 1):
            for image_path in _paragraph_images(
                paragraphs[paragraph_index],
                source_doc.part.related_parts,
                assets_dir,
                f"source_{paragraph_index}",
            ):
                paths.append(str(image_path.relative_to(source_dir)))
    raw_text = "\n".join(_paragraph_text(paragraph) for paragraph in paragraphs[start : end + 1]).strip()
    return packet_path, paths, raw_text


def extract_errata_items(job_id: str) -> None:
    _ensure_errata_task_column()
    with SessionLocal() as db:
        job = db.query(ErrataJob).filter(ErrataJob.job_id == job_id).first()
        if not job:
            return
        source_path = Path(job.source_path)
        try:
            anchors = _anchors(job.custom_anchors)
            doc = Document(str(source_path))
            paragraphs = _body_paragraphs(doc)
            texts = [_paragraph_text(paragraph) for paragraph in paragraphs]
            marker_indices = [idx for idx, text in enumerate(texts) if _is_marker(text, anchors)]
            if not marker_indices:
                raise ValueError(f"未找到勘误锚点。已尝试：{'、'.join(anchors)}")

            extracted_dir = source_path.parent / "evidence"
            extracted_dir.mkdir(parents=True, exist_ok=True)
            existing_items = {
                item.item_id: item
                for item in db.query(ErrataItem).filter(ErrataItem.job_id == job_id).all()
            }
            rebuilt_item_ids: set[str] = set()

            previous_marker = -1
            for item_index, marker_index in enumerate(marker_indices, start=1):
                block_start = next(
                    (
                        idx
                        for idx in range(marker_index - 1, previous_marker, -1)
                        if texts[idx].startswith(("原题：", "原始题目"))
                    ),
                    previous_marker + 1,
                )
                source_ref = next(
                    (
                        texts[idx]
                        for idx in range(block_start - 1, previous_marker, -1)
                        if texts[idx] and not _looks_like_next_item_metadata(texts[idx])
                    ),
                    "",
                )
                next_start = next(
                    (
                        idx
                        for idx in range(marker_index + 1, len(texts))
                        if texts[idx].startswith(("原题：", "原始题目"))
                    ),
                    len(texts),
                )
                existing_lines: list[str] = []
                existing_count = 0
                for idx in range(marker_index + 1, next_start):
                    if _looks_like_next_item_metadata(texts[idx]):
                        break
                    if texts[idx]:
                        existing_lines.append(texts[idx])
                        existing_count += 1

                evidence: list[str] = []
                for paragraph_index in range(block_start, marker_index):
                    for image_path in _paragraph_images(
                        paragraphs[paragraph_index],
                        doc.part.related_parts,
                        extracted_dir,
                        f"item_{item_index}_{paragraph_index}",
                    ):
                        evidence.append(str(image_path.relative_to(source_path.parent)))
                packet_path, material_paths, material_text = _build_material_packet(
                    source_path, doc, paragraphs, block_start, marker_index, item_index
                )
                question_end = next(
                    (
                        idx - 1
                        for idx in range(block_start, marker_index)
                        if texts[idx].startswith(("原答案：", "答案：", "修改意见：", "被反馈："))
                    ),
                    marker_index - 1,
                )
                question_end = max(block_start, question_end)
                _, question_material_paths, _ = _build_material_packet(
                    source_path,
                    doc,
                    paragraphs,
                    block_start,
                    question_end,
                    item_index,
                    "question_materials",
                )
                item_id = f"{job_id}_{item_index}"
                item = existing_items.pop(item_id, None)
                if not item:
                    item = ErrataItem(item_id=item_id, job_id=job_id, item_index=item_index)
                    db.add(item)
                else:
                    rebuilt_item_ids.add(item_id)
                item.source_ref = source_ref
                item.question_text = _field_text(texts, block_start, marker_index, "原题：", ("原答案：", "修改意见：", *anchors)) or _field_text(texts, block_start, marker_index, "原始题目", ("原答案：", "答案：", "修改意见：", *anchors))
                item.original_answer = _field_text(
                    texts,
                    block_start,
                    marker_index,
                    "原答案：",
                    ("修改意见：", "被反馈：", *anchors),
                )
                item.correction_opinion = _field_text(
                    texts, block_start, marker_index, "修改意见：", ("被反馈：", *anchors)
                ) or _field_text(texts, block_start, marker_index, "被反馈：", anchors)
                item.existing_content = "\n".join(existing_lines)
                item.existing_paragraph_count = existing_count
                item.evidence_json = json.dumps(evidence, ensure_ascii=False)
                item.source_start_index = block_start
                item.source_end_index = marker_index
                item.material_docx_path = str(packet_path)
                item.material_paths_json = json.dumps(material_paths, ensure_ascii=False)
                item.question_material_paths_json = json.dumps(
                    [
                        path
                        for path in question_material_paths
                        if Path(path).suffix.lower() in MODEL_IMAGE_SUFFIXES
                    ],
                    ensure_ascii=False,
                )
                item.material_text = material_text
                item.material_version = int(item.material_version or 0) + 1
                item.status = "pending"
                item.result_type = None
                item.review_status = "pending"
                item.review_feedback = None
                item.warnings_json = "[]"
                item.final_text_markup = None
                previous_marker = marker_index
            for item in existing_items.values():
                if not item.item_id.startswith(f"{job_id}_manual_"):
                    db.delete(item)
            job.state = "ready"
            job.error_msg = None
            db.flush()
            ensure_errata_tasks(db, job_id, create_missing=True)
            for item in db.query(ErrataItem).filter(ErrataItem.job_id == job_id).all():
                sync_errata_task(db, item, job, invalidate=item.item_id in rebuilt_item_ids)
            db.commit()
        except Exception as exc:
            job.state = "failed"
            job.error_msg = str(exc)
            db.commit()
            raise


def normalize_errata_evidence(job_id: str) -> int:
    """整理旧项目的自动证据，不重拆题、不覆盖人工编辑。"""
    with SessionLocal() as db:
        job = db.query(ErrataJob).filter(ErrataJob.job_id == job_id).first()
        if not job:
            raise ValueError("勘误任务不存在")
        source_path = Path(job.source_path)
        doc = Document(str(source_path))
        anchors = _anchors(job.custom_anchors)
        marker_indices = [
            index
            for index, paragraph in enumerate(_body_paragraphs(doc))
            if _is_marker(_paragraph_text(paragraph), anchors)
        ]
        page_images = _render_evidence(source_path, source_path.parent)
        marker_pages = _marker_pages(source_path, page_images, anchors)
        changed = 0
        items = (
            db.query(ErrataItem)
            .filter(ErrataItem.job_id == job_id)
            .order_by(ErrataItem.item_index)
            .all()
        )
        for item in items:
            # 手动新增题没有原 Word 锚点，保留用户手工归属的所有图片。
            if item.item_index > len(marker_indices):
                continue
            old_evidence = json.loads(item.evidence_json or "[]")
            embedded_evidence = [
                path
                for path in old_evidence
                if Path(path).parts and Path(path).parts[0] == "evidence"
            ]
            new_evidence = embedded_evidence
            marker_page = marker_pages[item.item_index - 1] if item.item_index <= len(marker_pages) else None
            if not new_evidence and marker_page and marker_page <= len(page_images):
                new_evidence = [
                    str(page_images[marker_page - 1].relative_to(source_path.parent))
                ]
            if new_evidence == old_evidence:
                continue
            item.evidence_json = json.dumps(new_evidence, ensure_ascii=False)
            item.review_status = "pending"
            if item.status == "confirmed":
                item.status = "pending"
            task = sync_errata_task(db, item, job, invalidate=True)
            if task:
                task.error_code = "图片证据已整理，请重新运行工作流。"
            changed += 1
        db.commit()
        return changed


def rebuild_errata_materials(job_id: str) -> None:
    """显式重建旧项目的材料包；会使已有勘误结果失效并要求重新运行。"""
    extract_errata_items(job_id)


def _image_data_url(path: Path) -> str:
    mime = mimetypes.guess_type(path.name)[0] or "image/png"
    return f"data:{mime};base64,{base64.b64encode(path.read_bytes()).decode('ascii')}"


def _set_errata_node_state(task_id: str, state: str, node_name: str) -> bool:
    with SessionLocal() as db:
        task = db.query(Task).filter(Task.task_id == task_id).first()
        if not task:
            return True
        if task.state in {"cancelled", "paused", "terminated", "abandoned"}:
            return True
        task.state = state
        task.current_node = node_name
        db.commit()
        return False


def _load_errata_payload(task_id: str, scope: Literal["question", "full"] = "question") -> tuple[dict, list[dict]]:
    with SessionLocal() as db:
        task = db.query(Task).filter(Task.task_id == task_id).first()
        item = (
            db.query(ErrataItem).filter(ErrataItem.item_id == task.source_item_id).first()
            if task and task.source_item_id
            else db.query(ErrataItem).filter(ErrataItem.task_id == task_id).first()
        )
        if not task or not item:
            raise ValueError("勘误 Task 没有关联 ErrataItem")
        job = db.query(ErrataJob).filter(ErrataJob.job_id == item.job_id).first()
        if not job:
            raise ValueError("勘误项目不存在")
        payload = item_to_dict(item)
        source_dir = Path(job.source_path).parent
        images: list[dict] = []
        material_paths = (
            payload.get("question_material_paths", [])
            if scope == "question"
            else [*payload.get("material_paths", []), *payload.get("evidence", [])]
        )
        material_paths = list(dict.fromkeys(material_paths))
        for relative_path in material_paths:
            evidence_path = (source_dir / relative_path).resolve()
            if (
                evidence_path.exists()
                and source_dir in evidence_path.parents
                and evidence_path.suffix.lower() in MODEL_IMAGE_SUFFIXES
            ):
                images.append(
                    {"type": "image_url", "image_url": {"url": _image_data_url(evidence_path)}}
                )
        text_available = bool(
            (payload.get("question_text") or "").strip()
            if scope == "question"
            else (payload.get("material_text") or "").strip()
        )
        if not images and not text_available:
            raise ValueError("题干文字和可用的 PNG/JPEG/WebP 图片均为空")
        return payload, images


def _record_errata_log(task_id: str, node_name: str, request: str, response: str) -> None:
    with SessionLocal() as db:
        db.add(
            AgentLog(
                task_id=task_id,
                node_name=node_name,
                request_payload=request,
                response_payload=response,
                cost_tokens=0,
            )
        )
        db.commit()


def _raw_errata_material(payload: dict, image_count: int) -> str:
    """勘误模型只接收题块材料包，避免旧字段拆分损坏原始语义。"""
    text = str(payload.get("material_text") or "").strip() or "（DOCX 未提取到可用文字，请以附图为准）"
    return (
        "以下是同一勘误题块的原始辅助文本，图片是该题块按原 DOCX 顺序渲染的视觉材料。"
        "不要依据任何旧字段名预设材料角色。\n\n"
        f"{text}\n\n附件图片：{image_count} 张。"
    )


async def errata_solver_node(state: AgentState) -> AgentState:
    """将受保护的勘误材料适配为标准 Solver 可读取的图片输入。"""
    import asyncio

    task_id = state["task_id"]
    if await asyncio.to_thread(solve_node_sync, task_id):
        return {**state, "status": "cancelled", "error_msg": "Task was manually cancelled."}
    try:
        payload, image_content = await asyncio.to_thread(_load_errata_payload, task_id, "question")
        image_urls = [entry["image_url"]["url"] for entry in image_content]
        question_text = (payload.get("question_text") or "").strip() or (
            f"题干文字未提取，请仅依据 {len(image_urls)} 张题干证据图片独立解题。"
        )
        result = await solve_image(
            image_urls,
            state.get("review_feedback"),
            (state.get("agent_configs") or {}).get("solver") or {},
            "errata_workflow",
            task_id,
            question_text,
        )
        draft = (result.get("draft") or "").strip()
        if not draft:
            raise ValueError("Solver 返回了空解题结果")
        await asyncio.to_thread(
            persist_task_artifact, task_id, "solver", draft,
            {"tokens": coerce_token_count(result.get("tokens"), 0), "workflow": "errata"},
            state.get("input_revision", 1),
        )
        return {
            **state,
            "status": "reviewing",
            "image_url": image_urls[0] if image_urls else state.get("image_url", ""),
            "image_urls": image_urls,
            "draft_solution": draft,
            "total_tokens": coerce_token_count(state.get("total_tokens"), 0) + coerce_token_count(result.get("tokens"), 0),
        }
    except Exception as exc:
        return {**state, "status": "failed", "failed_node": "solver", "error_msg": str(exc)}


async def errata_formatter_node(state: AgentState) -> AgentState:
    """复用普通 Formatter，并把完整正解保留给后续勘误裁决。"""
    result = await format_node(state)
    if result.get("status") in {"failed", "cancelled"}:
        return result
    solution = (result.get("final_result") or "").strip()
    if not solution:
        return {**result, "status": "failed", "failed_node": "formatter", "error_msg": "Formatter 返回了空正解。"}
    return {**result, "formatted_solution": solution, "final_result": None, "status": "reviewing"}


async def errata_review_node(state: AgentState) -> AgentState:
    import asyncio

    task_id = state["task_id"]
    if await asyncio.to_thread(_set_errata_node_state, task_id, "reviewing", "errata_adjudication"):
        return {**state, "status": "cancelled", "error_msg": "Task was manually stopped."}
    solution = (state.get("formatted_solution") or "").strip()
    if not solution:
        artifact = await asyncio.to_thread(
            latest_task_artifact, task_id, "formatter", state.get("input_revision", 1)
        )
        solution = (artifact.content if artifact else "").strip()
    if not solution:
        return {**state, "status": "failed", "failed_node": "errata_adjudication", "error_msg": "缺少完整解题结果"}
    try:
        payload, image_content = await asyncio.to_thread(_load_errata_payload, task_id, "full")
        prompt_bundle = get_prompt_bundle("errata_adjudication", "errata_workflow")
        context = (
            f"{_raw_errata_material(payload, len(image_content))}\n\n"
            f"原答案：\n{payload['original_answer']}\n\n"
            f"勘误意见：\n{payload['correction_opinion']}\n\n"
            f"独立完成的完整正解：\n{solution}"
        )
        user_prompt = (prompt_bundle.get("user") or "{errata_context}").replace("{errata_context}", context)
        config = {**(state.get("agent_configs") or {}).get("reviewer", {}), "streaming": False}
        timeout, max_retries = get_runtime_request_settings()
        result = await call_with_retry_and_fallback(
            create_llm_func=lambda runtime: get_llm(runtime).with_structured_output(ErrataAdjudication),
            messages=[
                SystemMessage(content=prompt_bundle.get("system") or "独立比较正解、原答案和勘误意见。"),
                HumanMessage(content=[{"type": "text", "text": user_prompt}, *image_content]),
            ],
            model_config=config,
            fallback_models=resolve_fallback_models("reviewer", config),
            timeout=timeout,
            max_retries=max_retries,
            task_id=task_id,
        )
        decision = result if isinstance(result, ErrataAdjudication) else ErrataAdjudication.model_validate(result)
        if decision.question_verdict != "incorrect":
            decision.question_errata = ""
        has_correction_opinion = bool((payload.get("correction_opinion") or "").strip())
        if not has_correction_opinion:
            decision.correction_opinion_verdict = "not_provided"
        allow_write = errata_allow_write(decision, has_correction_opinion)
        result_type = derive_errata_result_type(decision)
        artifact_content = decision.model_dump_json()
        await asyncio.to_thread(
            persist_task_artifact, task_id, "errata_adjudication", artifact_content,
            {"tokens": 0, "allow_write": allow_write, "result_type": result_type}, state.get("input_revision", 1),
        )
        await asyncio.to_thread(_record_errata_log, task_id, "errata_adjudication", context, artifact_content)
        with SessionLocal() as db:
            item = db.query(ErrataItem).filter(ErrataItem.task_id == task_id).first()
            if item:
                item.result_type = result_type
                item.warnings_json = json.dumps(decision.warnings, ensure_ascii=False)
                item.review_status = "passed" if allow_write else "needs_evidence"
                item.review_feedback = "；".join(decision.warnings)
                item.status = "generated" if allow_write else "insufficient_evidence"
                db.commit()
        return {
            **state,
            "formatted_solution": solution,
            "errata_decision": decision.model_dump(),
            "review_decision": "PASS" if allow_write else "INSUFFICIENT",
            "status": "formatting" if allow_write else "manual",
        }
    except Exception as exc:
        return {**state, "status": "failed", "failed_node": "errata_adjudication", "error_msg": str(exc)}


async def errata_format_node(state: AgentState) -> AgentState:
    import asyncio

    task_id = state["task_id"]
    if await asyncio.to_thread(_set_errata_node_state, task_id, "formatting", "word_composition"):
        return {**state, "status": "cancelled", "error_msg": "Task was manually stopped."}
    try:
        solution = (state.get("formatted_solution") or "").strip()
        if not solution:
            artifact = await asyncio.to_thread(latest_task_artifact, task_id, "formatter", state.get("input_revision", 1))
            solution = (artifact.content if artifact else "").strip()
        raw_decision = state.get("errata_decision")
        if not raw_decision:
            artifact = await asyncio.to_thread(latest_task_artifact, task_id, "errata_adjudication", state.get("input_revision", 1))
            raw_decision = json.loads(artifact.content) if artifact else None
        if not raw_decision:
            raise ValueError("缺少勘误裁决结果")
        decision = ErrataAdjudication.model_validate(raw_decision)
        if not errata_allow_write(
            decision, decision.correction_opinion_verdict != "not_provided"
        ):
            raise ValueError("勘误裁决不允许写入，任务必须转人工处理")
        markup = build_errata_word_content(decision, solution)
        result_type = derive_errata_result_type(decision)
        await asyncio.to_thread(
            persist_task_artifact, task_id, "word_composition", markup,
            {"format": "markdown", "result_type": result_type, "deterministic": True},
            state.get("input_revision", 1),
        )
        with SessionLocal() as db:
            item = db.query(ErrataItem).filter(ErrataItem.task_id == task_id).first()
            if item:
                item.final_text_markup = markup
                item.status = "completed"
                db.commit()
        return {**state, "status": "completed", "final_result": markup, "formatted_solution": solution, "errata_decision": decision.model_dump()}
    except Exception as exc:
        return {**state, "status": "failed", "failed_node": "word_composition", "error_msg": str(exc)}


def _mineru_field(block: str, label: str, next_labels: tuple[str, ...]) -> str:
    boundary = "|".join(map(re.escape, next_labels)) or "$"
    match = re.search(re.escape(label) + r"\s*(.*?)(?=" + boundary + r")", block, re.S)
    return match.group(1).strip() if match else ""


async def enrich_errata_job_with_mineru(
    job_id: str,
) -> None:
    with SessionLocal() as db:
        job = db.query(ErrataJob).filter(ErrataJob.job_id == job_id).first()
        if not job:
            return
        job.mineru_status = "parsing"
        source_path = job.source_path
        db.commit()
    try:
        markdown = await parse_local_file_with_mineru(source_path)
        blocks = re.split("|".join(re.escape(anchor) for anchor in _anchors(job.custom_anchors)), markdown)
        with SessionLocal() as db:
            job = db.query(ErrataJob).filter(ErrataJob.job_id == job_id).first()
            items = db.query(ErrataItem).filter(ErrataItem.job_id == job_id).order_by(ErrataItem.item_index).all()
            if not job:
                return
            job.mineru_markdown = markdown
            job.mineru_status = "ready"
            for index, item in enumerate(items):
                block = blocks[index] if index < len(blocks) else ""
                item.mineru_text = block
                if not (item.question_text or "").strip():
                    item.question_text = _mineru_field(block, "原题：", ("原答案：", "修改意见："))
                if not (item.original_answer or "").strip():
                    item.original_answer = _mineru_field(block, "原答案：", ("修改意见：",))
                if not (item.correction_opinion or "").strip():
                    item.correction_opinion = _mineru_field(block, "修改意见：", ())
                sync_errata_task(db, item, job, invalidate=True)
            db.commit()
    except Exception as exc:
        with SessionLocal() as db:
            job = db.query(ErrataJob).filter(ErrataJob.job_id == job_id).first()
            if job:
                job.mineru_status = "failed"
                job.error_msg = f"MinerU：{exc}"
                db.commit()


async def generate_errata_job(
    job_id: str,
    model_config: dict | None = None,
    reviewer_config: dict | None = None,
) -> None:
    with SessionLocal() as db:
        job = db.query(ErrataJob).filter(ErrataJob.job_id == job_id).first()
        if not job:
            return
        job.state = "generating"
        item_ids = [
            row[0]
            for row in db.query(ErrataItem.item_id)
            .filter(
                ErrataItem.job_id == job_id,
                Task.state.in_(["manual", "failed", "paused", "terminated", "abandoned"]),
            )
            .join(Task, Task.task_id == ErrataItem.task_id)
            .order_by(ErrataItem.item_index)
            .all()
        ]
        db.commit()
    for item_id in item_ids:
        await run_errata_task_by_item(item_id)
    with SessionLocal() as db:
        job = db.query(ErrataJob).filter(ErrataJob.job_id == job_id).first()
        if job:
            job.state = "reviewing"
            db.commit()


async def run_errata_task(
    task_id: str,
    model_config: dict | None = None,
    reviewer_config: dict | None = None,
    start_node: str = "solver",
    target_nodes: list[str] | None = None,
) -> None:
    del model_config, reviewer_config
    with SessionLocal() as db:
        task = db.query(Task).filter(Task.task_id == task_id).first()
        if not task or task.state in {"cancelled", "paused", "terminated", "abandoned"}:
            return
        task.state = "queued"
        task.error_code = None
        if start_node == "solver":
            task.retry_count = 0
            try:
                history = json.loads(task.history or "{}")
            except json.JSONDecodeError:
                history = {}
            for key in (
                "draft_solution",
                "formatted_solution",
                "errata_decision",
                "review_decision",
                "review_feedback",
                "failed_node",
            ):
                history.pop(key, None)
            task.history = json.dumps(history, ensure_ascii=False)
        db.commit()
    from app.main import run_agent_workflow_async

    await run_agent_workflow_async(
        task_id,
        start_node,
        target_nodes or ["solver", "reviewer", "formatter", "errata_adjudication", "word_composition"],
    )


async def run_errata_task_batch(task_ids: list[str]) -> None:
    """批量提交由统一 Task 信号量限流，避免 BackgroundTasks 顺序执行。"""
    import asyncio

    await asyncio.gather(
        *(run_errata_task(task_id) for task_id in task_ids),
        return_exceptions=True,
    )


async def run_errata_task_by_item(
    item_id: str,
    model_config: dict | None = None,
    reviewer_config: dict | None = None,
) -> None:
    with SessionLocal() as db:
        item = db.query(ErrataItem).filter(ErrataItem.item_id == item_id).first()
        task = db.query(Task).filter(Task.task_id == item.task_id).first() if item and item.task_id else None
        if not item or not task:
            return
        task_id = task.task_id
    await run_errata_task(task_id, model_config, reviewer_config)


_MARK_START = "ERRATAHIGHLIGHTSTART"
_MARK_END = "ERRATAHIGHLIGHTEND"


class _LegacyMarkParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self.in_mark = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.lower() != "mark" or attrs or self.in_mark:
            raise ValueError("历史高亮内容只允许精确的 <mark> 标签")
        self.in_mark = True
        self.parts.append(_MARK_START)

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() != "mark" or not self.in_mark:
            raise ValueError("历史 <mark> 标签未正确闭合")
        self.in_mark = False
        self.parts.append(_MARK_END)

    def handle_data(self, data: str) -> None:
        self.parts.append(data)

    def close(self) -> None:
        super().close()
        if self.in_mark:
            raise ValueError("历史 <mark> 标签未闭合")


def _prepare_legacy_markdown(value: str) -> tuple[str, bool]:
    text_value = (value or "").strip()
    if not re.search(r"</?mark(?:\s|>)", text_value, flags=re.I):
        return text_value, False
    protected_math: list[str] = []

    def protect_math(match: re.Match) -> str:
        protected_math.append(match.group(0))
        return f"ERRATAMATH{len(protected_math) - 1}TOKEN"

    protected_value = re.sub(r"\$[^$\r\n]+\$", protect_math, text_value)
    parser = _LegacyMarkParser()
    parser.feed(protected_value)
    parser.close()
    prepared = "".join(parser.parts)
    for index, math_text in enumerate(protected_math):
        prepared = prepared.replace(f"ERRATAMATH{index}TOKEN", math_text)
    return prepared, True


def _highlight_legacy_markers(doc: Document) -> None:
    highlighted = False
    for paragraph in _body_paragraphs(doc):
        active = False
        for run_element in list(paragraph._p.findall(qn("w:r"))):
            text_nodes = list(run_element.iter(qn("w:t")))
            run_text = "".join(node.text or "" for node in text_nodes)
            if not run_text:
                continue
            pieces = re.split(f"({_MARK_START}|{_MARK_END})", run_text)
            if len(pieces) == 1:
                if active:
                    run_properties = run_element.get_or_add_rPr()
                    highlight = OxmlElement("w:highlight")
                    highlight.set(qn("w:val"), "yellow")
                    run_properties.append(highlight)
                continue
            insert_at = paragraph._p.index(run_element)
            paragraph._p.remove(run_element)
            for piece in pieces:
                if piece == _MARK_START:
                    active = True
                    highlighted = True
                    continue
                if piece == _MARK_END:
                    active = False
                    continue
                if not piece:
                    continue
                clone = deepcopy(run_element)
                clone_text_nodes = list(clone.iter(qn("w:t")))
                for node in clone_text_nodes:
                    node.text = ""
                if clone_text_nodes:
                    clone_text_nodes[0].text = piece
                if active:
                    run_properties = clone.get_or_add_rPr()
                    highlight = OxmlElement("w:highlight")
                    highlight.set(qn("w:val"), "yellow")
                    run_properties.append(highlight)
                paragraph._p.insert(insert_at, clone)
                insert_at += 1
    if highlighted:
        return


def _pandoc_docx_blocks(markdown: str, reference_doc: Path, target_doc: Document) -> list[object]:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise ValueError("DOCX 导出失败：未找到 pandoc 命令")
    prepared, has_legacy_mark = _prepare_legacy_markdown(markdown)
    with tempfile.TemporaryDirectory(prefix="errata_pandoc_") as temp_dir:
        markdown_path = Path(temp_dir) / "answer.md"
        output_path = Path(temp_dir) / "answer.docx"
        markdown_path.write_text(prepared, encoding="utf-8")
        result = subprocess.run(
            [
                pandoc,
                "-f",
                "markdown+tex_math_dollars",
                str(markdown_path),
                "--reference-doc",
                str(reference_doc),
                "-o",
                str(output_path),
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        if result.returncode != 0 or not output_path.exists():
            detail = (result.stderr or result.stdout or "").strip()
            raise ValueError(f"DOCX 导出失败：Pandoc 转换失败{f'：{detail}' if detail else ''}")
        converted = Document(str(output_path))
        if has_legacy_mark:
            _highlight_legacy_markers(converted)
        try:
            converted_numbering_part = converted.part.part_related_by(RT.NUMBERING)
        except KeyError:
            converted_numbering_part = None
        if converted_numbering_part is not None:
            try:
                target_numbering_part = target_doc.part.part_related_by(RT.NUMBERING)
            except KeyError:
                target_numbering_part = NumberingPart(
                    converted_numbering_part.partname,
                    converted_numbering_part.content_type,
                    deepcopy(converted_numbering_part.element),
                    target_doc.part.package,
                )
                target_doc.part.relate_to(target_numbering_part, RT.NUMBERING)
            else:
                converted_numbering = converted_numbering_part.element
                target_numbering = target_numbering_part.element
                for tag_name, id_name in (("w:abstractNum", "w:abstractNumId"), ("w:num", "w:numId")):
                    known_ids = {
                        element.get(qn(id_name))
                        for element in target_numbering.findall(qn(tag_name))
                    }
                    for element in converted_numbering.findall(qn(tag_name)):
                        if element.get(qn(id_name)) not in known_ids:
                            target_numbering.append(deepcopy(element))
        return [
            deepcopy(element)
            for element in converted.element.body
            if element.tag != qn("w:sectPr")
        ]


def _remove_existing_after_anchor(anchor: Paragraph, count: int) -> None:
    next_element = anchor._p.getnext()
    removed = 0
    while next_element is not None and removed < count:
        following = next_element.getnext()
        if next_element.tag == qn("w:p"):
            next_element.getparent().remove(next_element)
            removed += 1
        next_element = following


def _insert_blocks_after(anchor_element, blocks: list[object]) -> None:
    current = anchor_element
    for block in blocks:
        current.addnext(block)
        current = block


def export_errata_job(job_id: str) -> Path:
    with SessionLocal() as db:
        job = db.query(ErrataJob).filter(ErrataJob.job_id == job_id).first()
        if not job:
            raise ValueError("勘误任务不存在")
        all_items = (
            db.query(ErrataItem)
            .filter(ErrataItem.job_id == job_id)
            .order_by(ErrataItem.item_index)
            .all()
        )
        if not all_items:
            raise ValueError("没有可导出的勘误题")
        task_map = {
            task.task_id: task
            for task in db.query(Task).filter(Task.task_id.in_([item.task_id for item in all_items if item.task_id])).all()
        }
        pending = [
            item for item in all_items
            if item.status != "confirmed"
            or not item.task_id
            or task_map.get(item.task_id) is None
            or task_map[item.task_id].state != "completed"
            or not (task_map[item.task_id].final_result or "").strip()
        ]
        if pending:
            raise ValueError(f"仍有 {len(pending)} 题未完成或未人工确认，不能导出")
        rows = (
            db.query(ErrataItem, Task)
            .join(Task, Task.task_id == ErrataItem.task_id)
            .filter(
                ErrataItem.job_id == job_id,
                ErrataItem.status == "confirmed",
                Task.state == "completed",
                Task.final_result.isnot(None),
            )
            .order_by(ErrataItem.item_index)
            .all()
        )
        if not rows:
            raise ValueError("至少完成一条勘误 Task 后才能导出")
        source_path = Path(job.source_path)
        output_path = source_path.parent / f"{source_path.stem}_已处理.docx"

        doc = Document(str(source_path))
        anchors_for_job = _anchors(job.custom_anchors)
        anchors = [
            paragraph
            for paragraph in _body_paragraphs(doc)
            if _is_marker(_paragraph_text(paragraph), anchors_for_job)
        ]
        for item, task in rows:
            final_markup = (task.final_result or "").strip()
            blocks = _pandoc_docx_blocks(final_markup, source_path, doc)
            if item.item_index > len(anchors):
                doc.add_paragraph(f"【手动新增勘误】{item.source_ref or f'题块 {item.item_index}'}")
                _insert_blocks_after(doc.paragraphs[-1]._p, blocks)
                continue
            anchor = anchors[item.item_index - 1]
            if item.replace_existing:
                _remove_existing_after_anchor(anchor, int(item.existing_paragraph_count or 0))
            _insert_blocks_after(anchor._p, blocks)

        doc.save(str(output_path))
        job.output_path = str(output_path)
        job.state = "completed"
        db.commit()
        return output_path
