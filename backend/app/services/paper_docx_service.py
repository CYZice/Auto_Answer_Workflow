import json
import os
import re
import shutil
import subprocess
import tempfile
import uuid
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.core.database import DEFAULT_DB_PATH, SessionLocal
from app.models.domain import PaperProject, PaperQuestion, Task
from app.services.markdown_parser import MarkdownParser
from app.services.mineru_ingestion import parse_local_file_with_mineru


PAPER_ROOT = Path(
    os.getenv("PAPER_DATA_DIR", str(DEFAULT_DB_PATH.parent / "papers"))
).resolve()
GROUP_RE = re.compile(r"^([一二三四五六七八九十百]+)[、.．]\s*(.+)$")
QUESTION_RE = re.compile(r"^(\d+)[、.．]\s*(.*)$")
ANSWER_BOUNDARY_RE = re.compile(r"参考答案|答案及解析|【解析】")


def _paragraph_text(paragraph) -> str:
    return "".join(
        node.text or "" for node in paragraph._p.iter() if node.tag in {qn("w:t"), qn("m:t")}
    ).strip()


def _paragraph_images(paragraph, document: Document, target_dir: Path) -> list[str]:
    paths: list[str] = []
    for blip in paragraph._p.iter(qn("a:blip")):
        rel_id = blip.get(qn("r:embed"))
        if not rel_id or rel_id not in document.part.related_parts:
            continue
        part = document.part.related_parts[rel_id]
        suffix = Path(str(part.partname)).suffix or ".png"
        name = f"image_{uuid.uuid4().hex[:10]}{suffix}"
        path = target_dir / name
        path.write_bytes(part.blob)
        paths.append(str(path))
    return paths


def extract_paper_questions(paper_id: str) -> list[PaperQuestion]:
    with SessionLocal() as db:
        project = db.query(PaperProject).filter(PaperProject.paper_id == paper_id).first()
        if not project:
            raise ValueError("试卷项目不存在")
        project.state = "extracting"
        db.commit()
        source_path = Path(project.source_path)

    document = Document(source_path)
    media_dir = source_path.parent / "media"
    media_dir.mkdir(parents=True, exist_ok=True)
    groups: list[dict] = []
    current_group = "未分组"
    group_index = 0
    current: dict | None = None

    for paragraph in document.paragraphs:
        text = _paragraph_text(paragraph)
        group_match = GROUP_RE.match(text)
        if group_match:
            name = group_match.group(2).strip()
            if ANSWER_BOUNDARY_RE.search(name):
                if current:
                    groups.append(current)
                    current = None
                break
            if current:
                groups.append(current)
                current = None
            group_index += 1
            current_group = name or f"第{group_index}题组"
            continue
        if ANSWER_BOUNDARY_RE.fullmatch(text) or (
            ANSWER_BOUNDARY_RE.search(text) and not QUESTION_RE.match(text)
        ):
            if current:
                groups.append(current)
            break

        question_match = QUESTION_RE.match(text)
        if question_match:
            if current:
                groups.append(current)
            if group_index == 0:
                group_index = 1
            number = question_match.group(1)
            body = question_match.group(2).strip()
            current = {
                "group_index": group_index,
                "group_name": current_group,
                "number": number,
                "parts": [body] if body else [],
                "images": _paragraph_images(paragraph, document, media_dir),
            }
            continue
        if current:
            if text:
                current["parts"].append(text)
            current["images"].extend(_paragraph_images(paragraph, document, media_dir))
    else:
        if current:
            groups.append(current)

    if not groups:
        with SessionLocal() as db:
            project = db.query(PaperProject).filter(PaperProject.paper_id == paper_id).first()
            project.state = "failed"
            project.error_msg = "未能识别题号。请确认文档题目以“1.”或“1、”开头。"
            db.commit()
        raise ValueError("未能从 DOCX 中识别题目")

    with SessionLocal() as db:
        db.query(PaperQuestion).filter(PaperQuestion.paper_id == paper_id).delete()
        items: list[PaperQuestion] = []
        for index, group in enumerate(groups, start=1):
            item = PaperQuestion(
                paper_id=paper_id,
                item_index=index,
                stable_key=f"{group['group_index']}.{group['number']}",
                group_name=group["group_name"],
                question_number=group["number"],
                question_text="\n".join(group["parts"]).strip(),
                image_paths_json=json.dumps(group["images"], ensure_ascii=False),
                enabled=True,
                state="pending",
            )
            db.add(item)
            items.append(item)
        project = db.query(PaperProject).filter(PaperProject.paper_id == paper_id).first()
        project.state = "ready"
        project.error_msg = None
        db.commit()
        for item in items:
            db.refresh(item)
        return items


def question_to_dict(item: PaperQuestion, task: Task | None = None) -> dict:
    return {
        "id": item.id,
        "paper_id": item.paper_id,
        "item_index": item.item_index,
        "stable_key": item.stable_key,
        "group_name": item.group_name,
        "question_number": item.question_number,
        "question_text": item.question_text,
        "image_paths": json.loads(item.image_paths_json or "[]"),
        "enabled": bool(item.enabled),
        "task_id": item.task_id,
        "state": task.state if task else item.state,
        "answer": (task.final_result or task.answer_preview or "") if task else "",
        "error_code": task.error_code if task else None,
    }


def project_to_dict(project: PaperProject, questions: list[dict]) -> dict:
    return {
        "paper_id": project.paper_id,
        "original_filename": project.original_filename,
        "state": project.state,
        "error_msg": project.error_msg,
        "mineru_status": project.mineru_status,
        "questions": questions,
    }


async def enrich_paper_project_with_mineru(
    paper_id: str,
) -> None:
    with SessionLocal() as db:
        project = db.query(PaperProject).filter(PaperProject.paper_id == paper_id).first()
        if not project:
            return
        project.mineru_status = "parsing"
        source_path = project.source_path
        db.commit()
    try:
        markdown = await parse_local_file_with_mineru(source_path)
        with SessionLocal() as db:
            project = db.query(PaperProject).filter(PaperProject.paper_id == paper_id).first()
            if project:
                project.mineru_markdown = markdown
                project.mineru_status = "ready"
                db.commit()
    except Exception as exc:
        with SessionLocal() as db:
            project = db.query(PaperProject).filter(PaperProject.paper_id == paper_id).first()
            if project:
                project.mineru_status = "failed"
                project.error_msg = f"MinerU：{exc}"
                db.commit()


def apply_mineru_paper_questions(paper_id: str) -> list[PaperQuestion]:
    with SessionLocal() as db:
        project = db.query(PaperProject).filter(PaperProject.paper_id == paper_id).first()
        if not project or not project.mineru_markdown:
            raise ValueError("MinerU 候选题目尚未就绪")
        old_items = (
            db.query(PaperQuestion)
            .filter(PaperQuestion.paper_id == paper_id)
            .order_by(PaperQuestion.item_index)
            .all()
        )
        if any(item.task_id for item in old_items):
            raise ValueError("已有题目开始解答，不能再替换为 MinerU 拆题结果")
        fallback_images = [item.image_paths_json or "[]" for item in old_items]
        questions = MarkdownParser().parse(project.mineru_markdown)
        if not questions:
            raise ValueError("MinerU Markdown 未识别出题号")
        db.query(PaperQuestion).filter(PaperQuestion.paper_id == paper_id).delete()
        group_names: list[str] = []
        items: list[PaperQuestion] = []
        for index, question in enumerate(questions, start=1):
            group_name = question.question_type or "未分类"
            if group_name not in group_names:
                group_names.append(group_name)
            group_index = group_names.index(group_name) + 1
            item = PaperQuestion(
                paper_id=paper_id,
                item_index=index,
                stable_key=f"{group_index}.{question.number}",
                group_name=group_name,
                question_number=str(question.number),
                question_text=question.content,
                image_paths_json=fallback_images[min(index - 1, len(fallback_images) - 1)] if fallback_images else "[]",
                enabled=True,
                state="pending",
            )
            db.add(item)
            items.append(item)
        project.mineru_status = "applied"
        project.state = "ready"
        db.commit()
        for item in items:
            db.refresh(item)
        return items


def export_paper_answers(paper_id: str) -> Path:
    with SessionLocal() as db:
        project = db.query(PaperProject).filter(PaperProject.paper_id == paper_id).first()
        if not project:
            raise ValueError("试卷项目不存在")
        questions = (
            db.query(PaperQuestion)
            .filter(PaperQuestion.paper_id == paper_id, PaperQuestion.enabled.is_(True))
            .order_by(PaperQuestion.item_index)
            .all()
        )
        task_ids = [item.task_id for item in questions if item.task_id]
        tasks = db.query(Task).filter(Task.task_id.in_(task_ids)).all() if task_ids else []
        task_map = {task.task_id: task for task in tasks}
        missing = [item.stable_key for item in questions if not item.task_id or not task_map.get(item.task_id) or not (task_map[item.task_id].final_result or task_map[item.task_id].answer_preview)]
        if missing:
            raise ValueError(f"以下题目尚无答案：{', '.join(missing)}")
        source_path = Path(project.source_path)

    answer_lines = ["# 参考答案", ""]
    last_group = None
    for item in questions:
        task = task_map[item.task_id]
        if item.group_name != last_group:
            answer_lines.extend([f"## {item.group_name}", ""])
            last_group = item.group_name
        answer = (task.answer_preview or task.final_result or "").strip()
        answer_lines.extend([f"**{item.question_number}、** {answer}", ""])

    document = Document(source_path)
    document.add_page_break()
    pandoc = shutil.which("pandoc")
    if pandoc:
        with tempfile.TemporaryDirectory() as temp_dir:
            markdown_path = Path(temp_dir) / "answers.md"
            appendix_path = Path(temp_dir) / "answers.docx"
            markdown_path.write_text("\n".join(answer_lines), encoding="utf-8")
            subprocess.run(
                [pandoc, str(markdown_path), "-o", str(appendix_path)],
                check=True,
                capture_output=True,
                timeout=120,
            )
            appendix = Document(appendix_path)
            for child in appendix.element.body:
                if child.tag != qn("w:sectPr"):
                    document.element.body.insert(-1, deepcopy(child))
    else:
        document.add_heading("参考答案", level=1)
        last_group = None
        for item in questions:
            task = task_map[item.task_id]
            if item.group_name != last_group:
                document.add_heading(item.group_name, level=2)
                last_group = item.group_name
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{item.question_number}、").bold = True
            paragraph.add_run((task.answer_preview or task.final_result or "").strip())

    output_path = source_path.with_name(f"{source_path.stem}_含参考答案.docx")
    document.save(output_path)
    with SessionLocal() as db:
        project = db.query(PaperProject).filter(PaperProject.paper_id == paper_id).first()
        project.output_path = str(output_path)
        project.state = "completed"
        db.commit()
    return output_path
