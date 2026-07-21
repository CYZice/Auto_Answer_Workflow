import base64
import json
import uuid
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from app.core.database import Base, SessionLocal, engine
from app.models.domain import ErrataItem, ErrataJob, PaperProject, PaperQuestion, Task, TaskArtifact
from app.services.errata_service import (
    MARKER,
    ErrataAdjudication,
    _load_errata_payload,
    _prepare_legacy_markdown,
    _raw_errata_material,
    build_errata_word_content,
    derive_errata_result_type,
    errata_allow_write,
    export_errata_job,
    extract_errata_items,
)
from app.services.paper_docx_service import apply_mineru_paper_questions, extract_paper_questions
from app.services.task_artifacts import persist_task_artifact


def test_errata_extracts_15_blocks_and_exports_mark_highlight(tmp_path: Path):
    Base.metadata.create_all(bind=engine)
    source = tmp_path / "errata.docx"
    document = Document()
    for index in range(1, 16):
        document.add_paragraph(f"题号 {index}")
        document.add_paragraph(f"原题：第 {index} 题题干 $x_{index}$")
        document.add_paragraph(f"原答案：第 {index} 题原答案")
        document.add_paragraph(f"修改意见：第 {index} 题修改意见")
        document.add_paragraph(MARKER)
    for relationship in list(document.part.rels.values()):
        if relationship.reltype == RT.NUMBERING:
            document.part.drop_rel(relationship.rId)
    document.save(source)
    with zipfile.ZipFile(source) as archive:
        assert "word/numbering.xml" not in archive.namelist()
    original_bytes = source.read_bytes()
    job_id = f"test_errata_{uuid.uuid4().hex}"
    with SessionLocal() as db:
        db.add(ErrataJob(job_id=job_id, original_filename=source.name, source_path=str(source), state="extracting"))
        db.commit()

    extract_errata_items(job_id)
    with SessionLocal() as db:
        items = db.query(ErrataItem).filter(ErrataItem.job_id == job_id).order_by(ErrataItem.item_index).all()
        assert len(items) == 15
        for index, item in enumerate(items, start=1):
            assert f"第 {index} 题题干" in item.question_text
            assert f"第 {index} 题原答案" in item.original_answer
            assert f"第 {index} 题修改意见" in item.correction_opinion
        items[0].status = "confirmed"
        items[0].result_type = "partial_fix"
        # 模拟人工在工作台修改题干和最终写入文本后的持久化状态。
        items[0].question_text = "人工修订后的题干"
        items[0].final_text_markup = "# 【解析】\n\n<mark>人工修正内容</mark>，由 $T_1<T_2$ 得 $x=\\boxed{2}$。"
        task = db.query(Task).filter(Task.task_id == items[0].task_id).one()
        task.state = "completed"
        task.final_result = items[0].final_text_markup
        for item in items[1:]:
            item.status = "confirmed"
            item.final_text_markup = "原答案正确。"
            item_task = db.query(Task).filter(Task.task_id == item.task_id).one()
            item_task.state = "completed"
            item_task.final_result = item.final_text_markup
        db.commit()

    output = export_errata_job(job_id)
    assert source.read_bytes() == original_bytes
    rendered = Document(output)
    highlighted = [run for paragraph in rendered.paragraphs for run in paragraph.runs if run.font.highlight_color == WD_COLOR_INDEX.YELLOW]
    assert [run.text for run in highlighted] == ["人工修正内容"]
    with zipfile.ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert "<m:oMath" in document_xml
    assert "# 【解析】" not in document_xml
    assert "\\boxed" not in document_xml

    with SessionLocal() as db:
        task_ids = [task_id for task_id, in db.query(ErrataItem.task_id).filter(ErrataItem.job_id == job_id).all() if task_id]
        db.query(ErrataItem).filter(ErrataItem.job_id == job_id).delete()
        if task_ids:
            db.query(Task).filter(Task.task_id.in_(task_ids)).delete(synchronize_session=False)
        db.query(ErrataJob).filter(ErrataJob.job_id == job_id).delete()
        db.commit()


def test_errata_material_packet_keeps_feedback_in_original_order(tmp_path: Path):
    Base.metadata.create_all(bind=engine)
    source = tmp_path / "packet.docx"
    image = tmp_path / "feedback.png"
    image.write_bytes(base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVQIHWP4z8DwHwAFgAI/ScLx8QAAAABJRU5ErkJggg=="))
    document = Document()
    document.add_paragraph("原题：混合材料题干")
    document.add_picture(str(image))
    document.add_paragraph("原答案：原解析内容")
    document.add_paragraph("被反馈：以下截图指出答案错误")
    document.add_picture(str(image))
    document.add_paragraph(MARKER)
    document.save(source)
    job_id = f"packet_{uuid.uuid4().hex}"
    with SessionLocal() as db:
        db.add(ErrataJob(job_id=job_id, original_filename=source.name, source_path=str(source)))
        db.commit()
    extract_errata_items(job_id)
    with SessionLocal() as db:
        item = db.query(ErrataItem).filter(ErrataItem.job_id == job_id).one()
        assert "混合材料题干" in item.material_text
        assert "被反馈" in item.material_text
        assert item.original_answer == "原解析内容"
        assert item.correction_opinion == "以下截图指出答案错误"
        assert item.material_docx_path and Path(item.material_docx_path).exists()
        assert item.material_paths_json != "[]"
        question_packet = source.parent / "question_materials" / "item_1" / "source.docx"
        question_text = "\n".join(paragraph.text for paragraph in Document(question_packet).paragraphs)
        assert "混合材料题干" in question_text
        assert "原解析内容" not in question_text
        assert "被反馈" not in question_text
        assert all(Path(path).suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"} for path in json.loads(item.question_material_paths_json or "[]"))
        payload = {"material_text": item.material_text, "question_text": "不应作为模型输入"}
        assert "不应作为模型输入" not in _raw_errata_material(payload, 1)
        item.result_type = "rewrite"
        item.warnings_json = json.dumps(["旧告警"], ensure_ascii=False)
        db.commit()
        extract_errata_items(job_id)
        db.expire_all()
        item = db.query(ErrataItem).filter(ErrataItem.job_id == job_id).one()
        assert item.result_type is None
        assert item.warnings_json == "[]"
        task_ids = [item.task_id] if item.task_id else []
        db.query(ErrataItem).filter(ErrataItem.job_id == job_id).delete()
        if task_ids:
            db.query(Task).filter(Task.task_id.in_(task_ids)).delete(synchronize_session=False)
        db.query(ErrataJob).filter(ErrataJob.job_id == job_id).delete()
        db.commit()


def test_errata_model_attachments_filter_emf(tmp_path: Path):
    Base.metadata.create_all(bind=engine)
    source = tmp_path / "source.docx"
    Document().save(source)
    evidence_dir = tmp_path / "evidence"
    evidence_dir.mkdir()
    (evidence_dir / "diagram.emf").write_bytes(b"emf")
    (evidence_dir / "page.png").write_bytes(base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVQIHWP4z8DwHwAFgAI/ScLx8QAAAABJRU5ErkJggg=="))
    suffix = uuid.uuid4().hex
    task_id = f"attachment_task_{suffix}"
    job_id = f"attachment_job_{suffix}"
    item_id = f"attachment_item_{suffix}"
    paths = json.dumps(["evidence/diagram.emf", "evidence/page.png"])
    with SessionLocal() as db:
        db.add(ErrataJob(job_id=job_id, original_filename=source.name, source_path=str(source)))
        db.add(Task(task_id=task_id, thread_id=task_id, image_url="", state="manual", retry_count=0, workflow_type="errata", source_kind="errata", source_id=job_id, source_item_id=item_id))
        db.add(ErrataItem(item_id=item_id, job_id=job_id, item_index=1, task_id=task_id, question_text="题干", original_answer="不应进入 Reviewer", material_text="完整材料", material_paths_json=paths, question_material_paths_json=paths))
        db.commit()

    _, question_images = _load_errata_payload(task_id, "question")
    _, full_images = _load_errata_payload(task_id, "full")
    assert len(question_images) == len(full_images) == 1
    assert all(entry["image_url"]["url"].startswith("data:image/png") for entry in question_images + full_images)

    with SessionLocal() as db:
        db.query(ErrataItem).filter(ErrataItem.item_id == item_id).delete()
        db.query(Task).filter(Task.task_id == task_id).delete()
        db.query(ErrataJob).filter(ErrataJob.job_id == job_id).delete()
        db.commit()


def test_errata_word_composition_copies_formatter_after_original_answer_error():
    solution = "【解析】由 $U=IR$ 得 $I=2A$。\n\n【结论】电流为 $2A$。"
    text = build_errata_word_content(
        ErrataAdjudication(
            standard_answer_verdict="correct",
            question_verdict="correct",
            original_answer_verdict="incorrect",
            correction_opinion_verdict="partial",
            errata_opinion="原答案漏写了<mark>单位</mark>。",
        ),
        solution,
    )
    assert text == solution


def test_errata_word_composition_uses_fixed_text_when_original_is_correct():
    text = build_errata_word_content(
        ErrataAdjudication(
            standard_answer_verdict="correct",
            question_verdict="correct",
            original_answer_verdict="correct",
            correction_opinion_verdict="incorrect",
            errata_opinion="计算结果与原答案一致。",
        ),
        "这段完整正解不应写入 Word。",
    )
    assert text == "原答案正确。"


def test_errata_word_composition_prefixes_question_errata():
    solution = "【解析】由 $T_1<T_2$ 得 $T_2>T_1$，故 $x=\\boxed{2}$。"
    text = build_errata_word_content(
        ErrataAdjudication(
            standard_answer_verdict="correct",
            question_verdict="incorrect",
            original_answer_verdict="incorrect",
            correction_opinion_verdict="correct",
            question_errata="题干中的方向应改为向右。",
        ),
        solution,
    )
    assert text == f"【题干勘误】题干中的方向应改为向右。\n\n{solution}"
    prepared, has_mark = _prepare_legacy_markdown(f"{solution}\n\n<mark>修正</mark>")
    assert has_mark is True
    assert "$T_1<T_2$" in prepared


def test_errata_verdict_fields_are_derived_consistently():
    cases = [
        ("correct", "correct", "correct", "not_provided", "correct", True),
        ("correct", "correct", "incorrect", "correct", "rewrite", True),
        ("correct", "incorrect", "incorrect", "correct", "question_errata", True),
        ("correct", "correct", "insufficient_evidence", "correct", "insufficient_evidence", False),
        ("incorrect", "correct", "incorrect", "correct", "rewrite", False),
    ]
    for standard, question, original, correction, result_type, allow_write in cases:
        decision = ErrataAdjudication(
            standard_answer_verdict=standard,
            question_verdict=question,
            original_answer_verdict=original,
            correction_opinion_verdict=correction,
            question_errata="题干修正" if question == "incorrect" else "",
        )
        assert derive_errata_result_type(decision) == result_type
        assert errata_allow_write(decision, correction != "not_provided") is allow_write


def test_errata_export_requires_every_item_to_be_confirmed(tmp_path: Path):
    Base.metadata.create_all(bind=engine)
    source = tmp_path / "confirm.docx"
    document = Document()
    for index in range(1, 3):
        document.add_paragraph(f"原题：第 {index} 题")
        document.add_paragraph(MARKER)
    document.save(source)
    job_id = f"confirm_{uuid.uuid4().hex}"
    with SessionLocal() as db:
        db.add(ErrataJob(job_id=job_id, original_filename=source.name, source_path=str(source)))
        db.commit()
    extract_errata_items(job_id)
    with SessionLocal() as db:
        items = db.query(ErrataItem).filter(ErrataItem.job_id == job_id).all()
        for item in items:
            task = db.query(Task).filter(Task.task_id == item.task_id).one()
            task.state = "completed"
            task.final_result = "勘误结果"
        items[0].status = "confirmed"
        db.commit()
    try:
        export_errata_job(job_id)
        assert False, "未确认题目不应允许导出"
    except ValueError as exc:
        assert "未完成或未人工确认" in str(exc)
    with SessionLocal() as db:
        task_ids = [task_id for task_id, in db.query(ErrataItem.task_id).filter(ErrataItem.job_id == job_id).all() if task_id]
        db.query(ErrataItem).filter(ErrataItem.job_id == job_id).delete()
        db.query(Task).filter(Task.task_id.in_(task_ids)).delete(synchronize_session=False)
        db.query(ErrataJob).filter(ErrataJob.job_id == job_id).delete()
        db.commit()


def test_paper_docx_extracts_two_groups_and_nine_questions(tmp_path: Path):
    Base.metadata.create_all(bind=engine)
    source = tmp_path / "paper.docx"
    document = Document()
    document.add_heading("电路期末考试", 0)
    document.add_paragraph("一、选择题")
    for index in range(1, 8):
        document.add_paragraph(f"{index}. 选择题 {index}")
    document.add_paragraph("二、计算题")
    for index in range(1, 3):
        document.add_paragraph(f"{index}. 计算题 {index}")
    document.add_page_break()
    document.add_paragraph("参考答案")
    document.add_paragraph("1. 不应被解析")
    document.save(source)
    paper_id = f"test_paper_{uuid.uuid4().hex}"
    with SessionLocal() as db:
        db.add(PaperProject(paper_id=paper_id, original_filename=source.name, source_path=str(source), state="extracting"))
        db.commit()

    extract_paper_questions(paper_id)
    with SessionLocal() as db:
        items = db.query(PaperQuestion).filter(PaperQuestion.paper_id == paper_id).order_by(PaperQuestion.item_index).all()
        assert len(items) == 9
        assert [item.stable_key for item in items] == ["1.1", "1.2", "1.3", "1.4", "1.5", "1.6", "1.7", "2.1", "2.2"]
        assert {item.group_name for item in items[:7]} == {"选择题"}
        assert {item.group_name for item in items[7:]} == {"计算题"}
        db.query(PaperQuestion).filter(PaperQuestion.paper_id == paper_id).delete()
        db.query(PaperProject).filter(PaperProject.paper_id == paper_id).delete()
        db.commit()


def test_artifact_keeps_workflow_start_revision():
    Base.metadata.create_all(bind=engine)
    task_id = f"test_task_{uuid.uuid4().hex}"
    with SessionLocal() as db:
        db.add(Task(task_id=task_id, thread_id=task_id, image_url="", state="manual", retry_count=0, input_revision=2))
        db.commit()
    persist_task_artifact(task_id, "solver", "old run output", {}, 1)
    with SessionLocal() as db:
        artifact = db.query(TaskArtifact).filter(TaskArtifact.task_id == task_id).one()
        assert artifact.input_revision == 1
        task = db.query(Task).filter(Task.task_id == task_id).one()
        assert task.input_revision == 2
        db.delete(artifact)
        db.delete(task)
        db.commit()


def test_mineru_markdown_can_replace_paper_split_without_losing_fallback_images(tmp_path: Path):
    Base.metadata.create_all(bind=engine)
    paper_id = f"test_mineru_{uuid.uuid4().hex}"
    source = tmp_path / "paper.docx"
    Document().save(source)
    markdown = """# 电路试卷
一、选择题
1. 第一题
2. 第二题
二、计算题
1. 第三题
参考答案
1. 不参与拆题
"""
    with SessionLocal() as db:
        db.add(PaperProject(paper_id=paper_id, original_filename=source.name, source_path=str(source), mineru_markdown=markdown, mineru_status="ready"))
        db.add(PaperQuestion(paper_id=paper_id, item_index=1, stable_key="1.1", group_name="旧题组", question_number="1", question_text="旧题", image_paths_json='["/tmp/image.png"]'))
        db.commit()

    items = apply_mineru_paper_questions(paper_id)
    assert [item.stable_key for item in items] == ["1.1", "1.2", "2.1"]
    assert items[0].group_name == "选择题"
    assert items[0].image_paths_json == '["/tmp/image.png"]'

    with SessionLocal() as db:
        db.query(PaperQuestion).filter(PaperQuestion.paper_id == paper_id).delete()
        db.query(PaperProject).filter(PaperProject.paper_id == paper_id).delete()
        db.commit()
